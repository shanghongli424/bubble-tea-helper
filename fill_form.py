#!/usr/bin/env python3
"""Generate a filled expert registration form as docx"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Title
title = doc.add_heading('附件一：专家申请入库登记表', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('专家基本信息', level=2)

fields = [
    ('姓名', '[待填写]'),
    ('性别', '[待填写]'),
    ('出生年月', '[待填写]'),
    ('身份证号码', '[待填写]'),
    ('工作单位', '[待填写]'),
    ('工作年限', '[待填写]'),
    ('职务', '[待填写]'),
    ('学历', '[待填写]'),
    ('职称', '[待填写]'),
    ('专家联系方式', ''),
    ('手 机', '[待填写]'),
    ('E-mail', '[待填写]'),
    ('评 审 专 业', '[待填写]'),
]

table = doc.add_table(rows=len(fields), cols=2)
table.style = 'Table Grid'
for i, (label, value) in enumerate(fields):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

doc.add_paragraph()

doc.add_heading('专家入库申请声明', level=2)
statement = doc.add_paragraph(
    '本人自愿申请加入中教集团专家库，接受中教集团对本人入库资格的评审，'
    '承诺以上所填信息均真实有效。本人如能通过评审并进入中教集团专家库，'
    '将自觉遵守以下工作守则：\n'
    '1、认真学习并执行《中教集团招标类采购管理 3.0》、《后勤部广东分部-采购管理办法》；\n'
    '2、履行职责，遵守纪律，严守秘密，廉洁自律，不自行与投标人进行可能影响评审结果的接触；\n'
    '3、不弄虚作假，不谋取私利，不泄露与评审活动有关的情况和资料；\n'
    '4、不收除审核评审工作劳务报酬以外的任何现金、有价证券和礼物，不收有关利害关系人的任何财物和好处；\n'
    '5、客观、公正、公平地参与评审工作，维护国家利益，维护招标、投标双方的合法权益；\n'
    '6、认真明确提出个人意见并对所提意见承担责任；\n'
    '7、如本人存在国家法律、法规或规章规定应该回避的情形时，将主动申请回避。'
)
statement.style = 'Normal'

doc.add_paragraph()
doc.add_paragraph('专家本人签字：_____________')
doc.add_paragraph(f'日    期：{datetime.now().strftime("%Y")}年{datetime.now().strftime("%m")}月{datetime.now().strftime("%d")}日')

doc.add_paragraph()
doc.add_paragraph('【审批栏位 - 由审批人填写】')
doc.add_paragraph('招标采购中心负责人审批：_____________')
doc.add_paragraph('后勤部广东分部负责人审批：_____________')

doc.add_paragraph()
doc.add_paragraph('说明：1.请随此申请登记表附专家本人身份证、学历证（如有）、职称证书（如有）复印件。')
doc.add_paragraph('2.专家本人信息变更时应及时通知我中教集团招标采购中心。')

# Page break for attachment 2
doc.add_page_break()

# Attachment 2
doc.add_heading('附件二：项目评标专家费支付申请表', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('项目评标专家费支付申请表', level=2)

fields2 = [
    ('学校', '[待填写]'),
    ('项目编号', '[待填写]'),
    ('项目名称', '[待填写]'),
    ('评审日期', '[待填写]'),
    ('评审专家', '共 _____ 人'),
    ('使用自有车辆', '□ 是  ■ 否（请选择）'),
    ('专家车辆登记', '[如有车辆请填写]'),
    ('评审费合计', '_____ 元'),
    ('车旅费合计', '_____ 元'),
    ('共计', '_____ 元'),
]

table2 = doc.add_table(rows=len(fields2), cols=2)
table2.style = 'Table Grid'
for i, (label, value) in enumerate(fields2):
    table2.rows[i].cells[0].text = label
    table2.rows[i].cells[1].text = value

doc.add_paragraph()
doc.add_paragraph('【审批栏位 - 由审批人填写】')
doc.add_paragraph('经办人：_____________')
doc.add_paragraph('招标采购中心负责人审批：_____________')
doc.add_paragraph('后勤部广东分部负责人审批：_____________')

doc.add_paragraph()
doc.add_paragraph('注：1、提交此表需一同附项目评审专家签到表复印件。')
doc.add_paragraph('2、上述填报项不够填写，可自行加行。')

output_path = '/Users/aaa/.openclaw/media/filled_form.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
